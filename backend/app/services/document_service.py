import json
import math
from typing import List, Tuple
from io import BytesIO
from sqlalchemy.orm import Session
from app.models.models import Document, DocumentChunk
from app.core.config import settings
import openai
import pypdf
import docx

_openai_client = None

def get_doc_openai_client():
    global _openai_client
    if _openai_client is None and settings.OPENAI_API_KEY:
        _openai_client = openai.AsyncOpenAI(api_key=settings.OPENAI_API_KEY)
    return _openai_client

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """Extract text from txt, pdf, docx"""
    text = ""
    ext = filename.split('.')[-1].lower()
    
    if ext == 'txt':
        text = file_content.decode('utf-8')
    elif ext == 'pdf':
        reader = pypdf.PdfReader(BytesIO(file_content))
        for page in reader.pages:
            text += page.extract_text() + "\n"
    elif ext == 'docx':
        doc = docx.Document(BytesIO(file_content))
        for para in doc.paragraphs:
            text += para.text + "\n"
    else:
        raise ValueError(f"Unsupported file format: {ext}")
        
    return text


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks"""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks


async def get_embedding(text: str) -> List[float]:
    """Get embedding from OpenAI"""
    client = get_doc_openai_client()
    if not client:
        # Fallback for dev mode without API key
        return [0.0] * 1536
        
    response = await client.embeddings.create(
        model="text-embedding-3-small",
        input=text
    )
    return response.data[0].embedding


def cosine_similarity(v1: List[float], v2: List[float]) -> float:
    """Calculate cosine similarity in pure Python"""
    if not v1 or not v2 or len(v1) != len(v2):
        return 0.0
    
    dot_product = sum(a * b for a, b in zip(v1, v2))
    norm_a = math.sqrt(sum(a * a for a in v1))
    norm_b = math.sqrt(sum(b * b for b in v2))
    
    if norm_a == 0.0 or norm_b == 0.0:
        return 0.0
        
    return dot_product / (norm_a * norm_b)


async def process_and_store_document(
    db: Session, 
    file_content: bytes, 
    filename: str
) -> Document:
    """Read file, chunk it, embed, and save to DB"""
    # 1. Extract text
    text = extract_text_from_file(file_content, filename)
    if not text.strip():
        raise ValueError("File is empty or could not extract text")
        
    # 2. Chunk text
    chunks = chunk_text(text)
    
    # 3. Create document record
    doc = Document(filename=filename)
    db.add(doc)
    db.flush() # get id
    
    # 4. Embed and save chunks
    for chunk in chunks:
        if not chunk.strip():
            continue
            
        emb = await get_embedding(chunk)
        
        doc_chunk = DocumentChunk(
            document_id=doc.id,
            chunk_text=chunk.strip(),
            embedding_json=json.dumps(emb)
        )
        db.add(doc_chunk)
        
    db.commit()
    db.refresh(doc)
    return doc


async def search_relevant_chunks(db: Session, query: str, top_k: int = 3) -> List[str]:
    """Search for most relevant chunks using pure Python cosine similarity"""
    client = get_doc_openai_client()
    if not client:
        return []
        
    # 1. Embed query
    query_emb = await get_embedding(query)
    if sum(query_emb) == 0.0:  # Fallback case
        return []
        
    # 2. Fetch all chunks
    # Note: This is fine for MVP/Small datasets (up to a few thousand chunks).
    # For large scale, use a proper Vector DB like Chroma/FAISS/pgvector.
    all_chunks = db.query(DocumentChunk).all()
    if not all_chunks:
        return []
        
    # 3. Calculate similarities
    similarities: List[Tuple[float, str]] = []
    for chunk in all_chunks:
        chunk_emb = json.loads(chunk.embedding_json)
        sim = cosine_similarity(query_emb, chunk_emb)
        similarities.append((sim, chunk.chunk_text))
        
    # 4. Sort and return top K
    similarities.sort(key=lambda x: x[0], reverse=True)
    
    # Filter out low confidence matches
    results = [text for sim, text in similarities[:top_k] if sim > 0.3]
    return results
